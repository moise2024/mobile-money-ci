"""
Génération du rapport technique PDF — Projet Mobile Money CI
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import subprocess

DOCX_PATH = os.path.join(os.path.dirname(__file__), 'data', 'output', 'Rapport_Projet_Mobile_Money_CI.docx')
PDF_PATH = DOCX_PATH.replace('.docx', '.pdf')
DASHBOARD_PATH = os.path.join(os.path.dirname(__file__), 'data', 'output', 'dashboard_mobile_money.png')

doc = Document()

# ─── Styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ─── PAGE DE GARDE ──────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph('')

titre = doc.add_paragraph()
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titre.add_run('DATA ENGINEERING\nProjet de Fin de Module')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(21, 101, 192)

doc.add_paragraph('')

sujet = doc.add_paragraph()
sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sujet.add_run("Analyse des flux Mobile Money en Côte d'Ivoire\nPipeline ETL & Schéma en étoile")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph('')
doc.add_paragraph('')

infos = [
    "Binôme : [Nom Prénom 1] & [Nom Prénom 2]",
    "Numéros étudiants : [XXXXXX] & [XXXXXX]",
    "Enseignant : GOUAH Tato Serge",
    "Date : Juin 2025",
    "GitHub : https://github.com/notre-projet/mobile-money-ci",
]
for info in infos:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(info)
    run.font.size = Pt(12)

doc.add_page_break()

# ─── 1. INTRODUCTION ────────────────────────────────────────────
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "La Côte d'Ivoire est l'un des marchés africains les plus dynamiques en matière de Mobile Money. "
    "Avec 4 opérateurs majeurs (MTN CI, Orange Money, Moov Africa, Wave) et plus de 20 millions "
    "d'utilisateurs, le secteur génère des millions de transactions chaque jour."
)
doc.add_paragraph(
    "Ce projet de fin de module nous place dans la peau d'un Data Engineer missionné par un agrégateur "
    "de paiements basé à Abidjan. Notre objectif : construire un pipeline de données complet, de la "
    "source brute jusqu'au tableau de bord analytique, en passant par une base de données cloud, "
    "l'orchestration et la documentation professionnelle."
)
doc.add_paragraph(
    "Problématique métier : comment transformer 100 000 transactions Mobile Money brutes en "
    "indicateurs actionnables pour la direction (parts de marché, tendances, performance des opérateurs) ?"
)

doc.add_heading('1.1 Objectifs du projet', level=2)
objectifs = [
    "Construire un pipeline ETL complet (Extract, Transform, Load) avec Python/Pandas",
    "Mettre en place un schéma en étoile analytique dans Supabase PostgreSQL",
    "Produire des requêtes SQL analytiques avec fonctions de fenêtre",
    "Créer un tableau de bord de 5 graphiques pour la direction",
    "Containeriser l'ensemble avec Docker et orchestrer avec Airflow",
]
for obj in objectifs:
    doc.add_paragraph(obj, style='List Bullet')

# ─── 2. ARCHITECTURE ────────────────────────────────────────────
doc.add_heading('2. Architecture du pipeline', level=1)
doc.add_paragraph(
    "Le pipeline suit l'architecture classique d'un projet Data Engineering moderne :"
)

arch_steps = [
    ("Source", "Fichier CSV (100 000 lignes, 12 colonnes)"),
    ("Extract", "Chargement avec Pandas, audit des types et valeurs manquantes"),
    ("Transform", "Nettoyage (NaN, aberrations), enrichissement (+8 colonnes), tests qualité"),
    ("Load", "Chargement dans Supabase PostgreSQL par lots de 5 000 lignes"),
    ("Modélisation", "Schéma en étoile : 4 dimensions + 1 table de faits"),
    ("Analyse", "Requêtes SQL avec JOIN, CTE, fonctions de fenêtre"),
    ("Visualisation", "Dashboard 5 graphiques avec Matplotlib"),
    ("Orchestration", "DAG Apache Airflow pour automatisation quotidienne"),
    ("Conteneurisation", "Docker + docker-compose pour reproductibilité"),
]
table = doc.add_table(rows=len(arch_steps), cols=2)
table.style = 'Light Shading Accent 1'
for i, (etape, desc) in enumerate(arch_steps):
    table.rows[i].cells[0].text = etape
    table.rows[i].cells[1].text = desc

doc.add_paragraph('')
doc.add_paragraph(
    "Le pipeline est conçu pour s'exécuter automatiquement chaque nuit à 23h00 via Airflow, "
    "avec 3 tentatives en cas d'échec et des alertes email configurées."
)

# ─── 3. SOURCES DE DONNÉES ──────────────────────────────────────
doc.add_heading('3. Sources de données', level=1)
doc.add_heading('3.1 Description du dataset', level=2)
doc.add_paragraph(
    "Le dataset fourni contient 100 000 transactions Mobile Money sur 6 mois "
    "(janvier à juin 2024). Chaque transaction est décrite par 12 colonnes :"
)

cols_desc = [
    ("id_transaction", "Identifiant unique de la transaction (MM0000001)"),
    ("date_heure", "Date et heure de la transaction (AAAA-MM-JJ HH:MM:SS)"),
    ("operateur", "Opérateur (MTN CI, Orange Money, Moov Africa, Wave)"),
    ("type_operation", "Type (Dépôt, Retrait, Transfert, Paiement marchand, Recharge téléphone)"),
    ("expediteur", "Numéro de l'émetteur"),
    ("beneficiaire", "Numéro du bénéficiaire"),
    ("montant_fcfa", "Montant en FCFA"),
    ("frais_fcfa", "Frais appliqués"),
    ("zone_expediteur", "Zone géographique de l'émetteur"),
    ("zone_beneficiaire", "Zone du bénéficiaire"),
    ("id_agent", "Identifiant de l'agent"),
    ("statut", "Statut (Succès, Échec, En attente)"),
]
table = doc.add_table(rows=len(cols_desc), cols=2)
table.style = 'Light Shading Accent 1'
for i, (col, desc) in enumerate(cols_desc):
    table.rows[i].cells[0].text = col
    table.rows[i].cells[1].text = desc

doc.add_heading('3.2 Audit initial', level=2)
doc.add_paragraph(
    "L'audit initial a révélé les problèmes suivants :\n"
    "• 2 000 valeurs manquantes sur frais_fcfa (2%)\n"
    "• 4 000 valeurs manquantes sur zone_beneficiaire (4%)\n"
    "• 3 000 valeurs manquantes sur id_agent (3%)\n"
    "• 200 montants aberrants (<= 0 FCFA)\n"
    "• Type date_heure incorrect (str au lieu de datetime)"
)

# ─── 4. ETL - EXTRACTION ────────────────────────────────────────
doc.add_heading('4. ETL — Extraction', level=1)
doc.add_paragraph(
    "L'extraction est réalisée avec pandas.read_csv(). Le fichier de 13.2 Mo est chargé "
    "en mémoire en 0.33 seconde. Les données sont ensuite auditées : types, valeurs "
    "manquantes, statistiques descriptives."
)
doc.add_paragraph(
    "Code clé : df_raw = pd.read_csv('transactions_mobile_money_100k.csv', encoding='utf-8')"
)

# ─── 5. ETL - TRANSFORMATION ────────────────────────────────────
doc.add_heading('5. ETL — Transformation', level=1)
doc.add_heading('5.1 Nettoyage', level=2)
doc.add_paragraph(
    "Les opérations de nettoyage appliquées :\n"
    "• Correction du type date_heure → datetime64\n"
    "• Uniformisation des chaînes vides en NaN\n"
    "• Imputation des frais manquants par 0 (recharges sans frais)\n"
    "• Imputation des zones manquantes par 'Zone inconnue'\n"
    "• Imputation des agents manquants par 'AGT-INCONNU'\n"
    "• Suppression des 200 lignes avec montant <= 0 FCFA"
)

doc.add_heading('5.2 Enrichissement', level=2)
doc.add_paragraph(
    "8 colonnes calculées ont été ajoutées pour enrichir l'analyse :"
)
enrich = [
    ("montant_net_fcfa", "Montant reçu par le bénéficiaire après déduction des frais"),
    ("taux_frais_pct", "Taux de frais en pourcentage du montant"),
    ("categorie_montant", "Tranche : Micro, Petit, Moyen, Gros"),
    ("tranche_horaire", "Période de la journée (Matin, Soirée, Nuit...)"),
    ("inter_ville", "1 si expéditeur et bénéficiaire sont dans des villes différentes"),
    ("heure", "Heure extraite de date_heure (0-23)"),
    ("jour_semaine", "Nom du jour en anglais"),
    ("mois", "Année-mois au format AAAA-MM"),
]
table = doc.add_table(rows=len(enrich), cols=2)
table.style = 'Light Shading Accent 1'
for i, (col, desc) in enumerate(enrich):
    table.rows[i].cells[0].text = col
    table.rows[i].cells[1].text = desc

doc.add_heading('5.3 Tests de qualité', level=2)
doc.add_paragraph(
    "5 règles de validation sont appliquées automatiquement :\n"
    "• Complétude : aucune valeur manquante sur les colonnes critiques\n"
    "• Unicité : aucun doublon sur id_transaction\n"
    "• Validité : montants strictement positifs (> 0 FCFA)\n"
    "• Conformité : statuts dans l'ensemble autorisé (Succès, Échec, En attente)\n"
    "• Typage : date_heure correctement convertie en datetime"
)

# ─── 6. ETL - CHARGEMENT ────────────────────────────────────────
doc.add_heading('6. ETL — Chargement Supabase', level=1)
doc.add_paragraph(
    "Supabase est une plateforme open-source fournissant une base PostgreSQL cloud gratuite "
    "(sans carte bancaire requise). Elle est idéale pour les projets data en Afrique."
)
doc.add_paragraph(
    "Le chargement s'effectue par lots (chunks) de 5 000 lignes via SQLAlchemy, avec "
    "la méthode 'multi' pour des insertions groupées plus rapides.\n\n"
    "Résultat : 99 800 lignes chargées dans la table 'transactions'.\n"
    "Tables d'agrégation créées : agg_operateur, agg_type_operation, agg_zone."
)

# ─── 7. MODÉLISATION ────────────────────────────────────────────
doc.add_heading('7. Modélisation — Schéma en étoile', level=1)
doc.add_paragraph(
    "Le schéma en étoile est conçu pour optimiser les requêtes analytiques. "
    "Il comprend 4 tables de dimensions et 1 table de faits :"
)

doc.add_heading('Tables de dimensions', level=2)
dims = [
    ("dim_operateur", "4 opérateurs avec type (Télécoms, Fintech)"),
    ("dim_zone", "11 zones avec région, type, population"),
    ("dim_type_operation", "5 types avec catégorie (Entrée/Sortie de fonds)"),
    ("dim_date", "182 jours (jan-juin 2024) avec année, mois, trimestre, jour, weekend"),
]
table = doc.add_table(rows=len(dims), cols=2)
table.style = 'Light Shading Accent 1'
for i, (dim, desc) in enumerate(dims):
    table.rows[i].cells[0].text = dim
    table.rows[i].cells[1].text = desc

doc.add_paragraph('')
doc.add_paragraph(
    "Table de faits : faits_transactions — 99 800 lignes avec clés étrangères "
    "pointant vers les 4 dimensions, plus les mesures (montant, frais, statut...)."
)

doc.add_paragraph(
    "L'ordre de peuplement est crucial : les dimensions doivent être chargées AVANT "
    "la table de faits, car les clés étrangères (foreign keys) référencent les IDs des dimensions."
)

# ─── 8. ANALYSES SQL ────────────────────────────────────────────
doc.add_heading('8. Analyses SQL', level=1)
doc.add_paragraph("4 requêtes analytiques ont été écrites et exécutées sur Supabase :")

doc.add_heading('Requête 1 : Volume par opérateur', level=2)
doc.add_paragraph(
    "Résultat : Les 4 opérateurs se partagent équitablement le marché (~25% chacun). "
    "Wave est légèrement en tête avec 3.73 milliards FCFA sur 6 mois."
)

doc.add_heading('Requête 2 : Tendances mensuelles', level=2)
doc.add_paragraph(
    "Résultat : Volume stable autour de 3 milliards FCFA par mois. Léger pic en février, "
    "légère baisse en juin. Taux de succès constant (~80%) sur toute la période."
)

doc.add_heading('Requête 3 : Jointure triple (zone × type × faits)', level=2)
doc.add_paragraph(
    "Résultat : Les dépôts sont l'opération dominante dans tous les quartiers d'Abidjan. "
    "Abidjan-Adjamé et Abidjan-Yopougon sont les zones les plus actives."
)

doc.add_heading('Requête 4 : CTE avec fonctions de fenêtre (RANK + LAG)', level=2)
doc.add_paragraph(
    "Résultat : Classement mensuel des opérateurs avec variation par rapport au mois précédent. "
    "L'utilisation de LAG() permet de mesurer la croissance mensuelle de chaque opérateur."
)

# ─── 9. TABLEAU DE BORD ─────────────────────────────────────────
doc.add_heading('9. Tableau de bord', level=1)
doc.add_paragraph(
    "Le tableau de bord contient 5 graphiques générés avec Matplotlib :"
)
graphs = [
    "Volume mensuel par opérateur (barres empilées)",
    "Parts de marché par volume (camembert)",
    "Volume par tranche horaire (courbe avec zones colorées)",
    "Taux de succès par opérateur (barres horizontales)",
    "Distribution des montants par catégorie (violin plot, échelle log)",
]
for g in graphs:
    doc.add_paragraph(g, style='List Bullet')

if os.path.exists(DASHBOARD_PATH):
    doc.add_paragraph('')
    doc.add_picture(DASHBOARD_PATH, width=Inches(6))
    doc.add_paragraph('Figure 1 : Tableau de bord Mobile Money CI — 5 graphiques')

doc.add_paragraph('')
doc.add_paragraph(
    "Interprétation : le marché est équilibré entre les 4 opérateurs. Les pics d'activité "
    "se situent le matin (8h-9h) et en soirée (18h-20h). La majorité des transactions "
    "sont de petit montant (Micro <= 5k FCFA)."
)

# ─── 10. ORCHESTRATION ──────────────────────────────────────────
doc.add_heading('10. Orchestration — Apache Airflow', level=1)
doc.add_paragraph(
    "Le DAG Airflow 'pipeline_mobile_money_ci' automatise l'exécution quotidienne du pipeline. "
    "Il se compose de 5 tâches enchaînées :"
)
tasks = [
    "extract_csv : Extraction CSV → Parquet brut",
    "clean_data : Nettoyage et enrichissement → Parquet propre",
    "load_supabase : Chargement dans Supabase PostgreSQL",
    "run_dbt : Exécution des modèles dbt (bonus)",
    "generate_report : Génération du rapport JSON de synthèse",
]
for t in tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    "Paramètres de résilience : 3 tentatives de reprise, délai de 5 minutes entre chaque essai, "
    "alerte email configurée en cas d'échec définitif. Le pipeline s'exécute automatiquement "
    "tous les jours à 23h00 (cron : 0 23 * * *)."
)

doc.add_heading('10.1 Conteneurisation Docker', level=2)
doc.add_paragraph(
    "L'ensemble du pipeline est containerisé avec Docker :\n"
    "• Dockerfile basé sur python:3.10-slim (~130 Mo)\n"
    "• docker-compose.yml avec 3 services : postgres, airflow-webserver, airflow-scheduler\n"
    "• Volumes partagés pour les DAGs et les données\n"
    "• Variables d'environnement pour les secrets (SUPABASE_URL)"
)

# ─── 11. CONCLUSION ─────────────────────────────────────────────
doc.add_heading('11. Conclusion', level=1)
doc.add_paragraph(
    "Ce projet nous a permis de mettre en pratique l'ensemble des compétences acquises "
    "durant le module Data Engineering :"
)
skills = [
    "Manipulation et nettoyage de données réelles avec Pandas (100k lignes)",
    "Construction d'un pipeline ETL complet avec tests de qualité automatisés",
    "Modélisation dimensionnelle (schéma en étoile) dans Supabase PostgreSQL",
    "Requêtes SQL analytiques avancées (CTE, fonctions de fenêtre)",
    "Visualisation des KPI avec Matplotlib",
    "Orchestration avec Apache Airflow (DAG, retries, alertes)",
    "Conteneurisation avec Docker et CI/CD avec GitHub Actions",
    "Introduction à dbt pour la transformation déclarative",
]
for s in skills:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    "Difficultés rencontrées : gestion des accents dans les statuts (Succès vs Succes), "
    "adaptation aux types datetime64 selon les versions de Pandas, optimisation du chargement "
    "par chunks dans Supabase (plan gratuit)."
)
doc.add_paragraph(
    "Améliorations possibles : ajout du streaming pour la détection de fraude en temps réel, "
    "passage à un Data Lakehouse (Delta Lake), déploiement sur un cluster Kubernetes, "
    "intégration d'une API REST pour l'accès aux données en temps réel."
)
doc.add_paragraph(
    "Remerciements à notre enseignant GOUAH Tato Serge pour la qualité du cours et "
    "l'encadrement tout au long du module."
)

# ─── SAUVEGARDE ─────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f'Rapport DOCX créé : {DOCX_PATH}')

# Conversion en PDF avec libreoffice
try:
    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', os.path.dirname(PDF_PATH), DOCX_PATH
    ], check=True, timeout=120)
    print(f'Rapport PDF créé : {PDF_PATH}')
except Exception as e:
    print(f'Conversion PDF impossible : {e}')
    print(f'Le fichier DOCX est disponible à la place.')
